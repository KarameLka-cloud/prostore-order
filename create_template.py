from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from seller import (
    SELLER_ACCOUNT,
    SELLER_BANK,
    SELLER_BIK,
    SELLER_CORR,
    SELLER_FULL,
    SELLER_INN,
    SELLER_NAME,
    SELLER_SIGN,
    SELLER_TAGLINE,
    resolve_logo_path,
    resolve_qr_maps_path,
    resolve_qr_telegram_path,
)

CONTENT_WIDTH_CM = 18.0


def _set_run_font(run, size=10, bold=False, name="Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, *, top=40, bottom=40, left=60, right=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _write_cell(
    cell,
    text: str,
    *,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size=9,
    clear=True,
) -> None:
    if clear:
        cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    _set_run_font(run, size=size, bold=bold)


def _add_para(doc, text: str, *, size=10, bold=False, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return p


def _set_table_borders(
    table, *, visible: bool, color: str = "AAAAAA", size: str = "4"
) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    val = "single" if visible else "nil"
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size if visible else "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color if visible else "FFFFFF")
        borders.append(node)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def _set_table_full_width(table, widths_cm: list[float]) -> None:
    assert abs(sum(widths_cm) - CONTENT_WIDTH_CM) < 0.05, widths_cm
    table.autofit = False
    table.allow_autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    for child in list(tbl_pr):
        if child.tag in (qn("w:tblW"), qn("w:tblLayout")):
            tbl_pr.remove(child)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    tbl_pr.append(tbl_w)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    widths = [Cm(w) for w in widths_cm]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def _write_stacked_cell(cell, lines: list[tuple[str, dict]]) -> None:
    cell.text = ""
    for i, (text, opts) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        _set_run_font(run, size=opts.get("size", 9),
                      bold=opts.get("bold", False))
    _set_cell_margins(cell)


def _add_header(doc: Document) -> None:
    header = doc.add_table(rows=1, cols=3)
    _set_table_borders(header, visible=False)
    _set_table_full_width(header, [4.2, 11.0, 2.8])

    logo_cell = header.rows[0].cells[0]
    text_cell = header.rows[0].cells[1]
    qr_cell = header.rows[0].cells[2]
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    logo_cell.text = ""
    logo_path = resolve_logo_path()

    if logo_path is not None:
        logo_p = logo_cell.paragraphs[0]
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after = Pt(0)
        run = logo_p.add_run()
        run.add_picture(str(logo_path), width=Cm(3.8))
    else:
        slot = logo_cell.add_table(rows=1, cols=1)
        slot.autofit = False
        _set_table_borders(slot, visible=True)
        slot_cell = slot.rows[0].cells[0]
        slot_cell.width = Cm(4.0)
        _set_cell_shading(slot_cell, "F5F5F5")
        _set_cell_margins(slot_cell, top=100, bottom=100, left=40, right=40)
        slot_cell.text = ""
        p = slot_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("ЛОГОТИП")
        _set_run_font(run, size=9, bold=True)
        hint = slot_cell.add_paragraph()
        hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hint.paragraph_format.space_before = Pt(2)
        hint.paragraph_format.space_after = Pt(0)
        run = hint.add_run("assets/logo.png")
        _set_run_font(run, size=7)
        run.font.color.rgb = RGBColor(120, 120, 120)

    text_cell.text = ""
    name_p = text_cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    run = name_p.add_run(SELLER_NAME)
    _set_run_font(run, size=12, bold=True)

    tag_p = text_cell.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(0)
    run = tag_p.add_run(SELLER_TAGLINE)
    _set_run_font(run, size=9)

    qr_cell.text = ""
    qr_path = resolve_qr_telegram_path()
    if qr_path is not None:
        qr_p = qr_cell.paragraphs[0]
        qr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_p.paragraph_format.space_before = Pt(0)
        qr_p.paragraph_format.space_after = Pt(0)
        run = qr_p.add_run()
        run.add_picture(str(qr_path), width=Cm(2.2))
        _set_cell_margins(qr_cell, top=0, bottom=0, left=40, right=0)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _merge_bank_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table, visible=True, color="000000", size="8")
    _set_table_full_width(table, [9.0, 2.4, 6.6])

    table.rows[0].cells[0].merge(table.rows[1].cells[0])
    table.rows[2].cells[1].merge(table.rows[3].cells[1])
    table.rows[2].cells[2].merge(table.rows[3].cells[2])

    bank_cell = table.rows[0].cells[0]
    bank_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _write_stacked_cell(
        bank_cell,
        [
            (SELLER_BANK, {"bold": True, "size": 9}),
            ("Банк получателя", {"size": 8}),
        ],
    )

    _write_cell(
        table.rows[0].cells[1],
        "БИК",
        size=8,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _write_cell(table.rows[0].cells[2], SELLER_BIK, size=9)
    _write_cell(
        table.rows[1].cells[1],
        "Кор. Счёт",
        size=8,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _write_cell(table.rows[1].cells[2], SELLER_CORR, size=9)

    for cell in (
        table.rows[0].cells[1],
        table.rows[0].cells[2],
        table.rows[1].cells[1],
        table.rows[1].cells[2],
    ):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell)

    _write_cell(table.rows[2].cells[0], f"ИНН {SELLER_INN}", size=9)
    table.rows[2].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(table.rows[2].cells[0])

    recipient = table.rows[3].cells[0]
    recipient.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _write_stacked_cell(
        recipient,
        [
            (SELLER_NAME, {"bold": True, "size": 9}),
            ("Получатель", {"size": 8}),
        ],
    )

    account_label = table.rows[2].cells[1]
    account_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _write_cell(account_label, "Счёт", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_margins(account_label, top=40, bottom=40, left=40, right=60)

    account_value = table.rows[2].cells[2]
    account_value.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _write_cell(account_value, SELLER_ACCOUNT, size=9)
    _set_cell_margins(account_value, top=40, bottom=40, left=60, right=60)


def _add_maps_banner(doc: Document) -> None:
    qr_path = resolve_qr_maps_path()
    if qr_path is None:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(qr_path), width=Cm(CONTENT_WIDTH_CM))


def build_order_docx(
    *,
    receipt_no: str,
    order_date: str,
    buyer: str,
    items: list[dict],
    total: str,
    total_words_line: str,
    path: Path,
) -> Path:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header(doc)

    _merge_bank_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    title = f"Товарный чек №{receipt_no} от {order_date}"
    _add_para(doc, title, size=14, bold=True, space_after=6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Продавец: ")
    _set_run_font(r, size=9, bold=True)
    r = p.add_run(SELLER_FULL)
    _set_run_font(r, size=9)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Покупатель: ")
    _set_run_font(r, size=9, bold=True)
    r = p.add_run(buyer or "Розничный покупатель")
    _set_run_font(r, size=9)

    headers = ["№", "Название товара или услуги",
               "Кол-во", "Ед. Изм.", "Цена", "НДС", "Сумма"]
    widths_cm = [1.0, 7.0, 1.5, 1.7, 2.4, 2.0, 2.4]
    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]

    table = doc.add_table(rows=1 + len(items), cols=7)
    table.style = "Table Grid"
    _set_table_full_width(table, widths_cm)
    widths = [Cm(w) for w in widths_cm]

    for i, (text, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        _write_cell(cell, text, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        _set_cell_shading(cell, "F0F0F0")
        _set_cell_margins(cell, top=30, bottom=30, left=40, right=40)

    keys = ["num", "name", "qty", "unit", "price", "vat", "sum"]
    for row_idx, item in enumerate(items, start=1):
        for col, (key, width, align) in enumerate(zip(keys, widths, aligns)):
            cell = table.rows[row_idx].cells[col]
            cell.width = width
            _write_cell(cell, item.get(key, ""), align=align, size=9)
            _set_cell_margins(cell, top=30, bottom=30, left=40, right=40)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    footer = doc.add_table(rows=1, cols=2)
    _set_table_borders(footer, visible=False)
    _set_table_full_width(footer, [13.5, 4.5])
    left = footer.rows[0].cells[0]
    right = footer.rows[0].cells[1]
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tc_pr = left._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)

    left.text = ""
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    r = lp.add_run(total_words_line)
    _set_run_font(r, size=9)
    _set_cell_margins(left, top=0, bottom=0, left=0, right=40)

    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)
    r = rp.add_run("Итог к оплате: ")
    _set_run_font(r, size=10)
    r = rp.add_run(f"{total} ₽")
    _set_run_font(r, size=10, bold=True)
    _set_cell_margins(right, top=0, bottom=0, left=40, right=0)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    sign = doc.add_paragraph()
    sign.paragraph_format.space_after = Pt(2)
    r = sign.add_run("Продавец:")
    _set_run_font(r, size=9, bold=True)

    for line in SELLER_SIGN.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _set_run_font(r, size=9)

    _add_maps_banner(doc)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
